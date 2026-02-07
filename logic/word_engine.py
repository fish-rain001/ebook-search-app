import os
import re
import glob
from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph

# =========================
# 路径
# =========================
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
BOOK_DIR = os.path.join(BASE_DIR, "data", "电子书")



# =========================
# 资料列表（新增）
# =========================
def list_collections():
    if not os.path.exists(BOOK_DIR):
        return []

    return sorted([
        d for d in os.listdir(BOOK_DIR)
        if os.path.isdir(os.path.join(BOOK_DIR, d))
    ])

# =========================
# 工具：遍历段落 + 表格
# =========================
def iter_block_items(parent):
    parent_elm = parent.element.body
    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('tbl'):
            yield Table(child, parent)

# =========================
# Word 样式判断（完全复刻你原代码）
# =========================
def is_heading1(para):
    return para.style.name in ['Heading 1', '标题 1', '标题1', '标题一', '标题', '一级标题']

def is_heading2(para):
    return para.style.name in ['Heading 2', '标题 2', '标题2', '标题二', '二级标题']

def is_normal_text(para):
    return para.style.name in ['Normal', '正文', 'Body Text', '正文文本']

# =========================
# 年份 / 期刊
# =========================
def list_years():
    if not os.path.exists(BOOK_DIR):
        return []
    return sorted(d.replace("年", "") for d in os.listdir(BOOK_DIR) if d.endswith("年"))

def list_issues(year):
    year_dir = os.path.join(BOOK_DIR, f"{year}年")
    issues = set()
    for f in os.listdir(year_dir):
        if f.endswith(".docx"):
            m = re.match(r"(第[一二三四五六七八九十\d]+期|第\d+号)", f)
            if m:
                issues.add(m.group(1))
    return sorted(list(issues))

def find_doc_path(year, issue):
    year_dir = os.path.join(BOOK_DIR, f"{year}年")
    patterns = [
        f"*{issue}*.docx",
        f"*{issue.replace('：','').replace(':','')}*.docx"
    ]
    for p in patterns:
        files = glob.glob(os.path.join(year_dir, p))
        if files:
            return files[0]
    return None

# =========================
# 专栏（标题1）
# =========================
def list_columns(doc_path):
    doc = Document(doc_path)
    columns = []
    skip = True

    for para in doc.paragraphs:
        text = para.text.strip()
        if re.match(r'0{60,}', text):
            skip = False
            continue
        if skip:
            continue
        if is_heading1(para) and text:
            columns.append(text)

    return columns

# =========================
# 主题（标题2）
# =========================
def list_topics(doc_path, column):
    doc = Document(doc_path)
    skip = True
    in_column = False
    topics = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if re.match(r'0{60,}', text):
            skip = False
            continue
        if skip:
            continue

        if is_heading1(para):
            norm_col = column.replace("：","").replace(":","")
            norm_text = text.replace("：","").replace(":","")
            if norm_col in norm_text:
                in_column = True
            elif in_column:
                break

        if in_column and is_heading2(para) and text:
            topics.append(text)

    return topics

# =========================
# 主题正文 + 表格
# =========================
def get_topic_content(doc_path, column, topic):
    doc = Document(doc_path)
    skip = True
    in_column = False
    in_topic = False
    content = []

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if re.match(r'0{60,}', text):
                skip = False
                continue
            if skip:
                continue

            if is_heading1(block):
                norm_col = column.replace("：","").replace(":","")
                norm_text = text.replace("：","").replace(":","")
                if norm_col in norm_text:
                    in_column = True
                elif in_column:
                    break

            if in_column and is_heading2(block):
                if text == topic:
                    in_topic = True
                    continue
                elif in_topic:
                    break

            if in_topic and is_normal_text(block) and text:
                content.append(text)

        elif isinstance(block, Table) and in_topic:
            rows = [[cell.text.strip() for cell in row.cells] for row in block.rows]
            content.append({"table": rows})

    return content

def full_text_search(doc_path, keyword):
    """
    完全复刻 Tkinter search_content 的文档扫描逻辑
    """
    if not keyword:
        return {
            "topics": [],
            "contents": [],
            "tables": []
        }

    doc = Document(doc_path)

    skip = True
    current_column = ""
    current_topic = ""

    topic_hits = []
    content_hits = []
    table_hits = []

    # ========= 段落扫描 =========
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name

        if re.match(r'0{60,}', text):
            skip = False
            continue
        if skip or not text:
            continue

        # 标题1：专栏
        if is_heading1(para):
            current_column = text
            if keyword.lower() in text.lower():
                topic_hits.append({
                    "column": current_column,
                    "topic": None,
                    "hit": text,
                    "type": "column"
                })
            continue

        # 标题2：主题
        if is_heading2(para):
            current_topic = text
            if keyword.lower() in text.lower():
                topic_hits.append({
                    "column": current_column,
                    "topic": current_topic,
                    "hit": text,
                    "type": "topic"
                })
            continue

        # 正文
        if is_normal_text(para):
            if keyword.lower() in text.lower():
                content_hits.append({
                    "column": current_column,
                    "topic": current_topic,
                    "content": text,
                    "type": "text"
                })

    # ========= 表格扫描 =========
    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if keyword.lower() in cell_text.lower():
                    rows = [[c.text.strip() for c in r.cells] for r in table.rows]
                    table_hits.append({
                        "column": current_column,
                        "topic": current_topic,
                        "content": rows,
                        "location": f"表格 第{r_idx+1}行 第{c_idx+1}列",
                        "type": "table"
                    })
                    break

    return {
        "topics": topic_hits,
        "contents": content_hits,
        "tables": table_hits
    }


